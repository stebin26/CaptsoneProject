"""Text extraction -- PDF, DOCX, and plain text to text with page metadata.

First stage of the RAG ingest pipeline. Page numbers are preserved through
extraction so an answer can later cite the exact page it came from, which is
what makes a grounded answer verifiable rather than merely plausible.
"""
# Text extraction — PDF / DOCX / TXT to plain text with page metadata.
# First stage of the RAG ingest pipeline; feeds the chunker.

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path

from ops_common.logging import get_logger

logger = get_logger(__name__)


class ExtractionError(RuntimeError):
    """Raised when a document exists but its text cannot be extracted."""

    pass


@dataclass
class ExtractedPage:
    """One page of extracted text with its page number."""
    page_number: int
    text: str


@dataclass
class ExtractedDocument:
    """A document after extraction: its pages plus identifying metadata."""
    filename: str
    file_type: str
    pages: list[ExtractedPage] = field(default_factory=list)

    @property
    def full_text(self) -> str:
        """Return every non-empty page joined into one string."""
        return "\n\n".join(p.text for p in self.pages if p.text.strip())

    @property
    def page_count(self) -> int:
        """Return the number of extracted pages."""
        return len(self.pages)


def detect_file_type(filename: str) -> str:
    """Determine a document's type from its filename.

    Markdown is reported as plain text, since it is extracted the same way.

    Args:
        filename: Name of the file being ingested.

    Returns:
        The normalized file type, or 'unknown' when there is no extension.
    """
    ext = Path(filename).suffix.lower().lstrip(".")
    if ext in ("pdf", "docx", "txt", "md"):
        return "txt" if ext == "md" else ext
    return ext or "unknown"


# ---------------------------------------------------------------------------
# Per-format extractors
# ---------------------------------------------------------------------------


def _extract_pdf(path: Path) -> list[ExtractedPage]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ExtractionError("pypdf is not installed; cannot read PDFs") from exc

    try:
        reader = PdfReader(str(path))
    except Exception as exc:
        # Encrypted, truncated, and non-PDF files renamed to .pdf all land here.
        logger.warning(
            "PDF could not be opened", extra={"file": path.name}, exc_info=True
        )
        raise ExtractionError(f"PDF could not be opened: {path.name} ({exc})") from exc

    pages: list[ExtractedPage] = []
    for i, page in enumerate(reader.pages):
        try:
            text = page.extract_text() or ""
        except Exception:  # noqa: BLE001 — a single bad page shouldn't kill the doc
            text = ""
        pages.append(ExtractedPage(page_number=i + 1, text=_clean(text)))
    return pages


def _extract_docx(path: Path) -> list[ExtractedPage]:
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:
        raise ExtractionError(
            "python-docx is not installed; cannot read DOCX files"
        ) from exc

    try:
        doc = DocxDocument(str(path))
    except Exception as exc:
        logger.warning(
            "DOCX could not be opened", extra={"file": path.name}, exc_info=True
        )
        raise ExtractionError(f"DOCX could not be opened: {path.name} ({exc})") from exc

    parts: list[str] = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    # docx has no intrinsic page breaks accessible here; treat as a single page.
    text = _clean("\n".join(parts))
    return [ExtractedPage(page_number=1, text=text)] if text else []


def _extract_txt(path: Path) -> list[ExtractedPage]:
    try:
        raw = path.read_text(encoding="utf-8", errors="ignore")
    except OSError as exc:
        logger.exception("Text file could not be read", extra={"file": path.name})
        raise ExtractionError(f"File could not be read: {path.name} ({exc})") from exc

    text = _clean(raw)
    return [ExtractedPage(page_number=1, text=text)] if text else []


# ---------------------------------------------------------------------------
# Cleaning + entry point
# ---------------------------------------------------------------------------


def _clean(text: str) -> str:
    # Normalize whitespace, drop null bytes, collapse excessive blank lines.
    text = text.replace("\x00", " ")
    lines = [ln.rstrip() for ln in text.splitlines()]
    cleaned: list[str] = []
    blank = 0
    for ln in lines:
        if ln.strip():
            cleaned.append(ln)
            blank = 0
        else:
            blank += 1
            if blank <= 1:
                cleaned.append("")
    return "\n".join(cleaned).strip()


def extract_document(
    path: str | Path, filename: str | None = None
) -> ExtractedDocument:
    """Extract text from a supported file.

    Args:
        path: Path to the stored file.
        filename: Original name to record; defaults to the path's own name.

    Returns:
        The extracted pages and their identifying metadata.

    Raises:
        ValueError: If the file type is not supported.
        ExtractionError: If the file is missing or its text cannot be read.
    """
    p = Path(path)
    name = filename or p.name
    ftype = detect_file_type(name)

    if not p.exists():
        raise ExtractionError(f"File not found: {p}")
    if not p.is_file():
        raise ExtractionError(f"Path is not a file: {p}")

    if ftype == "pdf":
        pages = _extract_pdf(p)
    elif ftype == "docx":
        pages = _extract_docx(p)
    elif ftype == "txt":
        pages = _extract_txt(p)
    else:
        raise ValueError(f"Unsupported file type: {ftype} ({name})")

    return ExtractedDocument(filename=name, file_type=ftype, pages=pages)
